"""
Сборка документа Freeze Form кодом (python-docx).
Текст берётся из редактируемого шаблона (dict), выравнивание — по левому краю.
"""
from __future__ import annotations
import io, json, copy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_FILE = Path(__file__).parent / "template_config.json"

# Поля «label -> ключ значения из данных» (порядок отображения сохраняется)
CLIENT_FIELDS = [
    ("field_full_name", "client_name"),
    ("field_plan", "selected_plan"),
    ("field_start", "start_date"),
    ("field_orig_exp", "orig_expiration"),
]
BREAK_FIELDS = [
    ("field_break_start", "break_start"),
    ("field_break_end", "break_end"),
    ("field_break_days", "break_days"),
]

DEFAULT_TEMPLATE = {
    "title1": "EXHIBIT {exhibit} TO CAREER SUPPORT SERVICES AGREEMENT",
    "title2": "OFFICIAL FREEZE FORM (PROGRAM PAUSE REQUEST)",
    "intro": ('This Official Freeze Form ("Form") is submitted pursuant to Section 16 '
              '(Break Policy) of the Career Support Services Agreement ("Agreement") '
              'between Go Offer Inc ("Provider") and the undersigned client ("Client").'),
    "client_header": "CLIENT INFORMATION",
    "field_full_name": "Full Name:",
    "field_plan": "Selected Plan:",
    "field_start": "Original Program Start Date:",
    "field_orig_exp": "Original Contract Expiration Date:",
    "break_header": "BREAK (FREEZE) DETAILS",
    "field_break_start": "Requested Break Start Date:",
    "field_break_end": "Requested Break End Date:",
    "field_break_days": "Total Duration of Break (in calendar days):",
    "note": ("Note: Standard Breaks may not exceed 3 weeks (21 days) without special "
             "written mutual agreement as per Section 16.1(b)."),
    "reason_label": "Reason for Break Request (Optional but recommended):",
    "timeline_header": "NEW CONTRACT TIMELINE (To be filled by Provider upon approval)",
    "field_adjusted": "Adjusted Contract Expiration Date:",
    "ack_header": "CLIENT ACKNOWLEDGEMENTS AND BINDING COMMITMENTS",
    "ack_intro": ("By signing and submitting this Form, the Client acknowledges and "
                  "agrees to the following:"),
    "ack_items": [
        {"title": "No Services During Break",
         "body": ("During the approved Break period, all active services (including "
                  "application submission, LinkedIn outreach, mentor calls, and curator "
                  "support) will be temporarily suspended.")},
        {"title": "Continued Obligation for Contingent Fee",
         "body": ("If the Client secures a Placement (receives and accepts a job offer) "
                  "during the Break period, the Client remains fully obligated to pay the "
                  "Contingent Fee as defined in Section 18.3 of the Agreement, provided the "
                  "interview process or communication was initiated according to the Offer "
                  "Attribution Criteria in Section 37.1.")},
        {"title": "No Infractions",
         "body": ("The Client will not receive Infractions under the Code of Conduct "
                  "(Section 24) for unresponsiveness during this approved Break period.")},
        {"title": "Approval Required",
         "body": ("This requested Break is not valid, and the Contract Expiration Date is "
                  "not officially adjusted, until this Form is approved and signed by an "
                  "authorized representative of Go Offer Inc.")},
    ],
    "provider_name": "Go Offer Inc",
    "provider_ein": "EIN: 93-4028120",
}


# ---------- персистентность шаблона ----------
def load_template() -> dict:
    if TEMPLATE_FILE.exists():
        try:
            cfg = json.loads(TEMPLATE_FILE.read_text(encoding="utf-8"))
            merged = copy.deepcopy(DEFAULT_TEMPLATE)
            merged.update(cfg)
            return merged
        except Exception:
            pass
    return copy.deepcopy(DEFAULT_TEMPLATE)


def save_template(cfg: dict) -> None:
    TEMPLATE_FILE.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")


def _sub(text: str, ctx: dict) -> str:
    for k, v in ctx.items():
        text = text.replace("{" + k + "}", str(v))
    return text


# ---------- помощники форматирования ----------
def _set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _bottom_rule(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom)
    pPr.append(pbdr)


# ---------- сборка ----------
def build_docx(cfg: dict, ctx: dict) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    def para(text="", *, bold=False, italic=False, center=False, after=0, runs=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(after)
        if runs:
            for t, b in runs:
                r = p.add_run(t); r.bold = b
        else:
            r = p.add_run(text); r.bold = bold; r.italic = italic
        return p

    # заголовки (по центру)
    para(_sub(cfg["title1"], ctx), bold=True, center=True, after=6)
    para(_sub(cfg["title2"], ctx), bold=True, center=True, after=10)

    # вступление
    para(_sub(cfg["intro"], ctx), after=12)

    # CLIENT INFORMATION
    para(cfg["client_header"], bold=True, after=2)
    for label_key, val_key in CLIENT_FIELDS:
        para(f"{cfg[label_key]} {ctx.get(val_key, '')}")
    doc.add_paragraph()

    # BREAK DETAILS
    para(cfg["break_header"], bold=True, after=2)
    for label_key, val_key in BREAK_FIELDS:
        para(f"{cfg[label_key]} {ctx.get(val_key, '')}")
    doc.add_paragraph()

    # note
    para(_sub(cfg["note"], ctx), italic=True, after=12)

    # reason
    para(cfg["reason_label"], after=4)
    para(ctx.get("reason", "N/A"), after=8)

    # divider
    _bottom_rule(doc.add_paragraph())
    doc.add_paragraph()

    # timeline
    para(cfg["timeline_header"], bold=True, after=2)
    para(f"{cfg['field_adjusted']} {ctx.get('adjusted_expiration', '')}", after=12)

    # acknowledgements
    para("", runs=[(cfg["ack_header"], True), (" " + _sub(cfg["ack_intro"], ctx), False)], after=8)
    for i, item in enumerate(cfg["ack_items"], 1):
        head = doc.add_paragraph()
        head.paragraph_format.left_indent = Inches(0.3)
        head.paragraph_format.space_after = Pt(0)
        r = head.add_run(f"{i}. "); r.bold = True
        r = head.add_run(_sub(item["title"], ctx)); r.bold = True
        b = doc.add_paragraph()
        b.paragraph_format.left_indent = Inches(0.55)
        b.paragraph_format.space_after = Pt(6)
        b.add_run(_sub(item["body"], ctx))
    doc.add_paragraph()

    # signature table
    rows = [
        ("Provider", "Client", True),
        (_sub(cfg["provider_name"], ctx), f"{cfg['field_full_name']} {ctx.get('client_name','')}", False),
        (_sub(cfg["provider_ein"], ctx), "Passport/Driving License:", False),
        ("Signature:", "Signature:", False),
        ("Date:", "Date:", False),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    _set_cell_borders(table)
    for ridx, (left, right, hdr) in enumerate(rows):
        for cidx, txt in enumerate((left, right)):
            cell = table.cell(ridx, cidx)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(txt)
            run.bold = hdr
            if ridx >= 3:  # место для подписи/даты
                cell.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    from freeze_core import FreezeData, docx_bytes_to_pdf_bytes
    from datetime import date
    cfg = load_template()
    d = FreezeData(client_name="Anna Gordeeva", selected_plan="Take All",
                   start_date=date(2026, 6, 16), orig_expiration=date(2026, 12, 16),
                   break_start=date(2026, 6, 16), break_days=10, exhibit="A", reason="")
    b = build_docx(cfg, d.context())
    Path("out").mkdir(exist_ok=True)
    Path("out/test.docx").write_bytes(b)
    Path("out/test.pdf").write_bytes(docx_bytes_to_pdf_bytes(b))
    print("docx:", len(b), "bytes; pdf written")
