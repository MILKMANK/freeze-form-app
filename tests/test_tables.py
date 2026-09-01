"""
Тесты поддержки таблиц данных в теле документа (engine.py).
Запуск: python3 tests/test_tables.py   (или pytest, если установлен)
Зависимости: python-docx (уже в requirements.txt).
"""
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
import engine as E


TABLE_TEXT = "\n".join([
    "# LAUNCH PACKAGE",
    "",
    "Hi {name},",
    "",
    "| Deliverable      | Delivered | Link   |",
    "| ---              | ---       | ---    |",
    "| Strategy session | {sdate}   | {slink} |",
    "| Master resume    | {rdate}   | {rlink} |",
    "",
    "The end.",
])

CTX = {"name": "John", "sdate": "Aug 01, 2026", "slink": "http://hub/s",
       "rdate": "Aug 02, 2026", "rlink": "http://hub/r"}


def _tables(text, ctx, with_signature=True):
    doc = Document(io.BytesIO(E.build_docx(text, ctx, ctx.get("name", ""),
                                           with_signature=with_signature)))
    return doc


def test_helpers_detect_table():
    assert E._is_table_row("| a | b |")
    assert not E._is_table_row("plain text")
    assert E._is_table_sep("| --- | --- |")
    assert E._is_table_sep("|---|:--:|")
    assert not E._is_table_sep("| a | b |")  # не разделитель — обычная строка данных


def test_split_cells_and_collect():
    assert E._split_cells("| a | b | c |") == ["a", "b", "c"]
    # неравное число колонок выравнивается к заголовку
    lines = TABLE_TEXT.split("\n")
    i = lines.index("| Deliverable      | Delivered | Link   |")
    header, rows, j = E._collect_table(lines, i)
    assert header == ["Deliverable", "Delivered", "Link"]
    assert len(rows) == 2
    assert all(len(r) == 3 for r in rows)
    assert j == i + 4  # заголовок + разделитель + 2 строки


def test_docx_has_data_table_with_substitution():
    doc = _tables(TABLE_TEXT, CTX, with_signature=False)
    assert len(doc.tables) == 1, "должна быть ровно одна таблица (данные, без подписей)"
    t = doc.tables[0]
    assert t.cell(0, 0).text == "Deliverable"
    assert t.cell(1, 0).text == "Strategy session"
    # {sdate}/{slink} подставились из ctx
    assert t.cell(1, 1).text == "Aug 01, 2026"
    assert t.cell(1, 2).text == "http://hub/s"
    assert t.cell(2, 1).text == "Aug 02, 2026"


def test_signature_table_still_added_alongside_data_table():
    doc = _tables(TABLE_TEXT, CTX, with_signature=True)
    # одна таблица данных + одна таблица подписей
    assert len(doc.tables) == 2
    sig = doc.tables[-1]
    assert "Provider" in sig.cell(0, 0).text
    assert "John" in sig.cell(1, 1).text


def test_single_pipe_line_is_not_a_table():
    """Строка на | без разделителя = обычный текст (обратная совместимость)."""
    text = "Just a | pipe | in text\nno table here"
    doc = _tables(text, {}, with_signature=False)
    assert len(doc.tables) == 0


def test_regression_freeze_and_ext_have_only_signature():
    for text in (E.FREEZE_TEXT, E.EXT_TEXT):
        doc = _tables(text, {"name": "X", "plan": "Y"}, with_signature=True)
        assert len(doc.tables) == 1, "во встроенных типах только таблица подписей"


def test_docx_hyperlink_is_clickable():
    text = "**[Raise an issue](https://t.me/go_offerIrina)**"
    doc = Document(io.BytesIO(E.build_docx(text, {}, "X", with_signature=False)))
    xml = doc.element.xml
    assert "hyperlink" in xml, "документ должен содержать w:hyperlink"
    targets = [r.target_ref for r in doc.part.rels.values() if "hyperlink" in r.reltype]
    assert "https://t.me/go_offerIrina" in targets
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Raise an issue" in full


def test_link_with_variable_substitution():
    text = "See [our hub]({link}) now."
    doc = Document(io.BytesIO(E.build_docx(text, {"link": "https://hub/x"}, "X", with_signature=False)))
    targets = [r.target_ref for r in doc.part.rels.values() if "hyperlink" in r.reltype]
    assert "https://hub/x" in targets, "переменная в url должна подставиться до парсинга ссылки"


def test_plain_text_without_link_unaffected():
    doc = Document(io.BytesIO(E.build_docx("Just **bold** text", {}, "X", with_signature=False)))
    assert "hyperlink" not in doc.element.xml
    assert "bold" in "\n".join(p.text for p in doc.paragraphs)


def _run_all():
    fns = [v for k, v in sorted(globals().items()) if k.startswith("test_") and callable(v)]
    failed = 0
    for fn in fns:
        try:
            fn()
            print(f"  PASS  {fn.__name__}")
        except Exception as e:
            failed += 1
            print(f"  FAIL  {fn.__name__}: {type(e).__name__}: {e}")
    print(f"\n{len(fns) - failed}/{len(fns)} passed")
    return failed


if __name__ == "__main__":
    sys.exit(1 if _run_all() else 0)
