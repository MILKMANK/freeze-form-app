"""
Движок Go Offer Docs.
- Расчёты дат и контекста для типов документов.
- Рендер текста с разметкой (**жирный**, *курсив*, # H1, ## H2, {переменные}) в .docx.
- Таблица подписей в конце документа (настоящая таблица).
- Конвертация .docx -> .pdf через LibreOffice.
"""
from __future__ import annotations
import io, re, subprocess, shutil, tempfile
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATE_FMT = "%B %d, %Y"
MONTHS = ["January","February","March","April","May","June","July","August",
          "September","October","November","December"]


# ------------- даты -------------
def add_months(d: date, n: int) -> date:
    m = d.month - 1 + n
    y = d.year + m // 12
    m = m % 12 + 1
    # последний день месяца
    if m == 12:
        last = 31
    else:
        last = (date(y, m + 1, 1) - timedelta(days=1)).day
    return date(y, m, min(d.day, last))


def fmt(d) -> str:
    return d.strftime(DATE_FMT) if d else "—"


def add_duration(d: date, n: int, unit: str) -> date:
    """unit: 'months' или 'days'."""
    return add_months(d, n) if unit == "months" else d + timedelta(days=n)


# типы-длительности и их единица
DURATION_UNIT = {"dur_days": "days", "dur_months": "months"}


# =================== ФОРМУЛЫ (Excel-подобные, без eval) ===================
class FormulaError(Exception):
    pass


def refs_in(expr: str):
    """Токены переменных, на которые ссылается формула."""
    return re.findall(r"\{(\w+)\}", expr or "")


_TOKEN_RE = re.compile(r"""\s*(
    \{\w+\}            |
    \d+\.?\d*%         |
    \d+\.?\d*          |
    "[^"]*"            |
    >=|<=|<>           |
    [-+*/(),;<>=&]     |
    [A-Za-z_][A-Za-z0-9_]*
)""", re.X)


def _tokenize(expr: str):
    return [m.group(1) for m in _TOKEN_RE.finditer(expr or "")]


def _is_num(v):
    return isinstance(v, (int, float)) and not isinstance(v, bool)


def _to_text(v):
    if isinstance(v, str):
        return v
    f = float(v)
    return str(int(f)) if f == int(f) else (f"{f:.4f}".rstrip("0").rstrip("."))


def _num(v, ctx="число"):
    if isinstance(v, str):
        raise FormulaError(f"ожидалось {ctx}, а не текст")
    return float(v)


def _edate_serial(serial, months):
    d = date.fromordinal(int(round(serial)))
    return add_months(d, int(round(months))).toordinal()


def _call(name, a):
    n = name
    def num(i): return _num(a[i])
    if n == "TODAY":
        return float(date.today().toordinal())
    if n == "DATE":
        return float(date(int(num(0)), int(num(1)), int(num(2))).toordinal())
    if n == "EDATE":
        return float(_edate_serial(num(0), num(1)))
    if n == "ADDDAYS":
        return num(0) + num(1)
    if n == "DAYS":
        return num(0) - num(1)
    if n in ("YEAR", "MONTH", "DAY"):
        d = date.fromordinal(int(round(num(0))))
        return float({"YEAR": d.year, "MONTH": d.month, "DAY": d.day}[n])
    if n == "ROUND":
        return float(round(num(0), int(num(1)) if len(a) > 1 else 0))
    if n == "ROUNDUP":
        import math
        k = 10 ** (int(num(1)) if len(a) > 1 else 0)
        return float(math.ceil(num(0) * k) / k)
    if n == "ROUNDDOWN":
        import math
        k = 10 ** (int(num(1)) if len(a) > 1 else 0)
        return float(math.floor(num(0) * k) / k)
    if n == "INT":
        return float(int(num(0)))
    if n == "ABS":
        return abs(num(0))
    if n == "MIN":
        return min(_num(x) for x in a)
    if n == "MAX":
        return max(_num(x) for x in a)
    if n == "SUM":
        return sum(_num(x) for x in a)
    if n == "IF":
        return a[1] if _num(a[0]) != 0 else (a[2] if len(a) > 2 else 0.0)
    if n == "AND":
        return 1.0 if all(_num(x) != 0 for x in a) else 0.0
    if n == "OR":
        return 1.0 if any(_num(x) != 0 for x in a) else 0.0
    if n == "NOT":
        return 1.0 if _num(a[0]) == 0 else 0.0
    if n == "UPPER":
        return _to_text(a[0]).upper()
    if n == "LOWER":
        return _to_text(a[0]).lower()
    if n == "LEN":
        return float(len(_to_text(a[0])))
    if n == "CONCAT":
        return "".join(_to_text(x) for x in a)
    if n == "TEXT":
        return _text_fn(a[0], _to_text(a[1]) if len(a) > 1 else "")
    raise FormulaError(f"нет функции {name}()")


_MONTHS_EN = ["January", "February", "March", "April", "May", "June", "July",
              "August", "September", "October", "November", "December"]


def _text_fn(val, f):
    if any(c in f for c in "YyMmDd"):
        d = date.fromordinal(int(round(float(val))))
        out = f
        out = out.replace("YYYY", f"{d.year:04d}").replace("MMMM", _MONTHS_EN[d.month - 1])
        out = out.replace("YY", f"{d.year % 100:02d}")
        out = out.replace("MM", f"{d.month:02d}").replace("DD", f"{d.day:02d}")
        return out
    if "." in f:
        dec = len(f.split(".")[1])
        return f"{float(val):.{dec}f}"
    if f:
        return str(int(round(float(val))))
    return _to_text(val)


class _Parser:
    def __init__(self, toks, env):
        self.t, self.i, self.env = toks, 0, env

    def peek(self):
        return self.t[self.i] if self.i < len(self.t) else None

    def nxt(self):
        tk = self.t[self.i]; self.i += 1; return tk

    def expect(self, x):
        if self.peek() != x:
            raise FormulaError(f'ожидалось "{x}"')
        self.i += 1

    def parse(self):
        v = self.expr()
        if self.i < len(self.t):
            raise FormulaError(f"лишнее: {self.peek()}")
        return v

    def expr(self):
        a = self.cmp()
        while self.peek() == "&":
            self.nxt(); a = _to_text(a) + _to_text(self.cmp())
        return a

    def cmp(self):
        a = self.add(); op = self.peek()
        if op in (">", "<", ">=", "<=", "=", "<>"):
            self.nxt(); b = self.add()
            try:
                res = {">": a > b, "<": a < b, ">=": a >= b, "<=": a <= b,
                       "=": a == b, "<>": a != b}[op]
            except TypeError:
                raise FormulaError("нельзя сравнить текст и число")
            return 1.0 if res else 0.0
        return a

    def add(self):
        a = self.mul()
        while self.peek() in ("+", "-"):
            op = self.nxt(); b = self.mul()
            a = _num(a) + _num(b) if op == "+" else _num(a) - _num(b)
        return a

    def mul(self):
        a = self.un()
        while self.peek() in ("*", "/"):
            op = self.nxt(); b = self.un()
            if op == "*":
                a = _num(a) * _num(b)
            else:
                if _num(b) == 0:
                    raise FormulaError("деление на ноль")
                a = _num(a) / _num(b)
        return a

    def un(self):
        if self.peek() == "-":
            self.nxt(); return -_num(self.un())
        return self.prim()

    def prim(self):
        x = self.peek()
        if x is None:
            raise FormulaError("пустое выражение")
        if x == "(":
            self.nxt(); e = self.expr(); self.expect(")"); return e
        if x[0] == "{":
            self.nxt(); tok = x[1:-1]
            if tok not in self.env:
                raise FormulaError(f"нет переменной {{{tok}}}")
            val = self.env[tok]
            if val is None:
                raise FormulaError(f"{{{tok}}} не заполнена")
            return val
        if x[0] == '"':
            self.nxt(); return x[1:-1]
        if x.endswith("%"):
            self.nxt(); return float(x[:-1]) / 100.0
        if x[0].isdigit():
            self.nxt(); return float(x)
        if re.match(r"^[A-Za-z_]", x):
            self.nxt(); self.expect("(")
            args = []
            if self.peek() != ")":
                args.append(self.expr())
                while self.peek() in (",", ";"):
                    self.nxt(); args.append(self.expr())
            self.expect(")")
            return _call(x.upper(), args)
        raise FormulaError(f"непонятно: {x}")


def evaluate_formula(expr: str, env: dict):
    return _Parser(_tokenize(expr), env).parse()


def _fmt_out(out, val):
    if out == "date":
        return fmt(date.fromordinal(int(round(float(val)))))
    if out == "number":
        f = float(val)
        return str(int(f)) if f == int(f) else (f"{f:.2f}".rstrip("0").rstrip("."))
    return _to_text(val)


def resolve_derived(defs: list, env: dict, vtypes: dict) -> dict:
    """Считает формулы и calc_date в порядке зависимостей. Мутирует env, возвращает ctx."""
    ctx = {}
    pending = list(defs)
    progress = True
    while pending and progress:
        progress = False; still = []
        for v in pending:
            tok = v["token"]
            if v["type"] == "calc_date":
                refs = [r for r in (v.get("base"), v.get("dur")) if r]
            else:
                refs = refs_in(v.get("expr", ""))
            if all(r in env for r in refs):
                try:
                    if v["type"] == "calc_date":
                        base = env.get(v.get("base")); durv = env.get(v.get("dur")) or 0
                        if base is None:
                            raise FormulaError("база не заполнена")
                        unit = DURATION_UNIT.get(vtypes.get(v.get("dur")), "days")
                        ser = _edate_serial(base, durv) if unit == "months" else base + durv
                        env[tok] = float(ser); ctx[tok] = fmt(date.fromordinal(int(round(ser))))
                    else:
                        val = evaluate_formula(v.get("expr", ""), env)
                        out = v.get("out", "number")
                        env[tok] = val
                        ctx[tok] = _fmt_out(out, val)
                except FormulaError as e:
                    env[tok] = None; ctx[tok] = "ошибка: " + str(e)
                except Exception:
                    env[tok] = None; ctx[tok] = "ошибка в формуле"
                progress = True
            else:
                still.append(v)
        pending = still
    for v in pending:  # циклы / нерешённые ссылки
        env[v["token"]] = None; ctx[v["token"]] = "—"
    return ctx


def compute_fields(defs: list, form: dict) -> dict:
    """
    Универсальный расчёт. type in
      text|number|date|dur_days|dur_months|calc_date|formula.
    Даты в окружении хранятся как серийный номер дня (ordinal).
    """
    env, ctx, vtypes, derived = {}, {}, {}, []
    for v in defs:
        tok, tp = v["token"], v["type"]
        vtypes[tok] = tp
        if tp == "date":
            val = form.get(tok)
            env[tok] = val.toordinal() if val else None
            ctx[tok] = fmt(val) if val else "—"
        elif tp in ("number", "dur_days", "dur_months"):
            val = form.get(tok)
            n = 0 if val in (None, "") else int(val)
            env[tok] = float(n); ctx[tok] = str(n)
        elif tp == "text":
            val = form.get(tok)
            env[tok] = "" if val in (None, "") else str(val); ctx[tok] = env[tok]
        elif tp in ("formula", "calc_date"):
            derived.append(v)
    ctx.update(resolve_derived(derived, env, vtypes))
    return {"err": None, "ctx": ctx, "raw": env}


# обратная совместимость
def compute_custom(fields: list, form: dict) -> dict:
    return compute_fields(fields, form)


# ------------- контекст по типам -------------
def compute_freeze(f: dict) -> dict:
    bs, st = f["break_start"], f["start_date"]
    oe = f["orig_exp"]
    if f.get("mode", "days") == "days":
        bd = max(1, int(f.get("break_days") or 1))
        bend = bs + timedelta(days=bd - 1)
        err = None
    else:
        bend = f.get("break_end")
        err = "Дата окончания паузы раньше даты начала." if (bend and bend < bs) else None
        bd = (bend - bs).days + 1 if bend else 0
    adj = oe + timedelta(days=bd) if (oe and bd) else None
    return {"err": err, "ctx": {
        "exhibit": f.get("exhibit", "A"),
        "name": f.get("client_name", ""),
        "plan": f.get("selected_plan", ""),
        "start_date": fmt(st),
        "orig_expiration": fmt(oe),
        "break_start": fmt(bs),
        "break_end": fmt(bend),
        "break_days": str(bd or "—"),
        "reason": (f.get("reason") or "").strip() or "N/A",
        "adjusted_expiration": fmt(adj),
    }, "raw": {
        "start_date": st, "orig_expiration": oe, "break_start": bs, "break_end": bend,
        "break_days": bd, "adjusted_expiration": adj,
    }}


def compute_extension(f: dict) -> dict:
    cur = f["current_exp"]
    m = max(1, int(f.get("ext_months") or 1))
    ne = add_months(cur, m) if cur else None
    return {"err": None, "ctx": {
        "exhibit": f.get("exhibit", "A"),
        "name": f.get("client_name", ""),
        "plan": f.get("selected_plan", ""),
        "start_date": fmt(f["start_date"]),
        "current_expiration": fmt(cur),
        "extension_months": str(m),
        "new_expiration": fmt(ne),
    }, "raw": {
        "start_date": f["start_date"], "current_expiration": cur,
        "extension_months": m, "new_expiration": ne,
    }}


# ------------- дефолтные шаблоны -------------
FREEZE_TEXT = "\n".join([
    "# EXHIBIT {exhibit} TO CAREER SUPPORT SERVICES AGREEMENT",
    "# OFFICIAL FREEZE FORM (PROGRAM PAUSE REQUEST)",
    "",
    'This Official Freeze Form ("Form") is submitted pursuant to Section 16 (Break Policy) '
    'of the Career Support Services Agreement ("Agreement") between Go Offer Inc ("Provider") '
    'and the undersigned client ("Client").',
    "",
    "## CLIENT INFORMATION",
    "Full Name: {name}",
    "Selected Plan: {plan}",
    "Original Program Start Date: {start_date}",
    "Original Contract Expiration Date: {orig_expiration}",
    "",
    "## BREAK (FREEZE) DETAILS",
    "Requested Break Start Date: {break_start}",
    "Requested Break End Date: {break_end}",
    "Total Duration of Break (in calendar days): {break_days}",
    "",
    "*Note: Standard Breaks may not exceed 3 weeks (21 days) without special written mutual "
    "agreement as per Section 16.1(b).*",
    "",
    "Reason for Break Request (Optional but recommended):",
    "{reason}",
    "",
    "## NEW CONTRACT TIMELINE (To be filled by Provider upon approval)",
    "Adjusted Contract Expiration Date: {adjusted_expiration}",
    "",
    "## CLIENT ACKNOWLEDGEMENTS AND BINDING COMMITMENTS",
    "By signing and submitting this Form, the Client acknowledges and agrees to the following:",
    "",
    "**1. No Services During Break**",
    "During the approved Break period, all active services (including application submission, "
    "LinkedIn outreach, mentor calls, and curator support) will be temporarily suspended.",
    "",
    "**2. Continued Obligation for Contingent Fee**",
    "If the Client secures a Placement (receives and accepts a job offer) during the Break period, "
    "the Client remains fully obligated to pay the Contingent Fee as defined in Section 18.3 of the "
    "Agreement, provided the interview process or communication was initiated according to the Offer "
    "Attribution Criteria in Section 37.1.",
    "",
    "**3. No Infractions**",
    "The Client will not receive Infractions under the Code of Conduct (Section 24) for "
    "unresponsiveness during this approved Break period.",
    "",
    "**4. Approval Required**",
    "This requested Break is not valid, and the Contract Expiration Date is not officially adjusted, "
    "until this Form is approved and signed by an authorized representative of Go Offer Inc.",
])

EXT_TEXT = "\n".join([
    "# EXHIBIT {exhibit} TO CAREER SUPPORT SERVICES AGREEMENT",
    "# OFFICIAL EXTENSION OF PROGRAM TERM",
    "",
    'This Extension of Program Term ("Extension") is entered into between Go Offer Inc ("Provider") '
    'and the undersigned client ("Client") to extend the term of the Career Support Services '
    'Agreement ("Agreement").',
    "",
    "## CLIENT INFORMATION",
    "Full Name: {name}",
    "Selected Plan: {plan}",
    "Original Program Start Date: {start_date}",
    "Current Contract Expiration Date: {current_expiration}",
    "",
    "## EXTENSION DETAILS",
    "Extension Length (in months): {extension_months}",
    "New Contract Expiration Date: {new_expiration}",
    "",
    "## CLIENT ACKNOWLEDGEMENTS AND BINDING COMMITMENTS",
    "By signing and submitting this Extension, the Client acknowledges and agrees to the following:",
    "",
    "**1. Continued Services**",
    "All services under the Agreement continue through the new Contract Expiration Date on the same "
    "terms and conditions.",
    "",
    "**2. Fees and Obligations**",
    "Any fees, contingent obligations, or commitments defined in the Agreement remain in full effect "
    "for the extended term.",
    "",
    "**3. Approval Required**",
    "This Extension is not valid, and the Contract Expiration Date is not officially adjusted, until "
    "this Form is approved and signed by an authorized representative of Go Offer Inc.",
])

FREEZE_VARS = [
    ("exhibit", "буква доп. соглашения (A, B, C…)"), ("name", "имя клиента"),
    ("plan", "выбранный план"), ("start_date", "дата создания Хаба"),
    ("orig_expiration", "изначальная дата окончания"), ("break_start", "дата старта паузы"),
    ("break_end", "дата окончания паузы (расчёт)"), ("break_days", "длительность паузы в днях"),
    ("reason", "причина (или N/A)"), ("adjusted_expiration", "новая дата окончания (расчёт)"),
]
EXT_VARS = [
    ("exhibit", "буква доп. соглашения"), ("name", "имя клиента"), ("plan", "выбранный план"),
    ("start_date", "дата создания Хаба"), ("current_expiration", "текущая дата окончания"),
    ("extension_months", "на сколько месяцев продление"), ("new_expiration", "новая дата окончания (расчёт)"),
]

PROVIDER_DEFAULT = {"name": "Go Offer Inc", "ein": "EIN: 93-4028120"}


# ------------- разметка -> docx -------------
def _sub(text: str, ctx: dict) -> str:
    return re.sub(r"\{(\w+)\}", lambda m: str(ctx.get(m.group(1), m.group(0))), text)


def _split(text: str, marker: str):
    parts = text.split(marker)
    return [(p, i % 2 == 1) for i, p in enumerate(parts)]


def _parse_runs(line: str):
    runs = []
    for bt, bold in _split(line, "**"):
        for it, ital in _split(bt, "*"):
            if it != "":
                runs.append((it, bold, ital))
    return runs


def _add_runs(p, line: str, ctx: dict, base_bold=False, size=None):
    for seg, bold, ital in _parse_runs(line):
        r = p.add_run(_sub(seg, ctx))
        r.bold = bold or base_bold
        r.italic = ital
        if size:
            r.font.size = Pt(size)


def _set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


# Таблица подписей одинакова во всех документах.
SIG_PROVIDER_NAME = "Go Offer Inc"
SIG_PROVIDER_EIN = "EIN: 93-4028120"


def _signature_table(doc, client_name: str):
    # (left, right, header?, кол-во пустых строк под заполнение)
    rows = [
        ("Provider", "Client", True, 0),
        (SIG_PROVIDER_NAME, f"Full Name: {client_name}", False, 2),
        (SIG_PROVIDER_EIN, "Passport/Driving License:", False, 2),
        ("Signature:", "Signature:", False, 4),
        ("Date:", "Date:", False, 2),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    _set_table_borders(table)
    for ri, (left, right, hdr, blanks) in enumerate(rows):
        for ci, txt in enumerate((left, right)):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = cell.paragraphs[0].add_run(txt)
            run.bold = hdr
            for _ in range(blanks):
                cell.add_paragraph()
    return table


def build_docx(text: str, ctx: dict, client_name: str, with_signature: bool = True) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    for raw in (text or "").split("\n"):
        line = raw.rstrip("\r")
        if line.startswith("## "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            _add_runs(p, line[3:], ctx, base_bold=True)
        elif line.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_runs(p, line[2:], ctx, base_bold=True, size=14)
        elif line.strip() == "":
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_runs(p, line, ctx)

    if with_signature:
        doc.add_paragraph()
        _signature_table(doc, client_name)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


# ------------- pdf -------------
def _find_soffice():
    for n in ("soffice", "libreoffice"):
        p = shutil.which(n)
        if p:
            return p
    for p in ("/Applications/LibreOffice.app/Contents/MacOS/soffice",
              r"C:\Program Files\LibreOffice\program\soffice.exe"):
        if Path(p).exists():
            return p
    return None


def docx_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    so = _find_soffice()
    if not so:
        raise RuntimeError("LibreOffice не найден — PDF недоступен.")
    d = Path(tempfile.mkdtemp())
    src = d / "doc.docx"; src.write_bytes(docx_bytes)
    subprocess.run([so, "--headless", "-env:UserInstallation=file:///tmp/lo_go_profile",
                    "--convert-to", "pdf", "--outdir", str(d), str(src)],
                   check=True, capture_output=True)
    pdf = d / "doc.pdf"
    if not pdf.exists():
        raise RuntimeError("PDF не создан.")
    return pdf.read_bytes()


def safe_name(s: str) -> str:
    return re.sub(r"[^\w\-]+", "_", s or "client").strip("_") or "client"
